# -*- coding: utf-8 -*-
"""Build the AI CUP 2026 VeriPromiseESG competition report by FILLING the official
blank template (ESG 永續承諾驗證競賽_報告空白範本.docx), so the layout, fonts,
margins and the author-contact table exactly match the organiser's required format.

Run:  python report/build_report.py
Out:  report/TEAM_10505_ESG永續承諾驗證競賽.docx   (then export to PDF in Word)
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

CN_FONT, EN_FONT = "標楷體", "Times New Roman"
HERE = Path(__file__).resolve().parent
TEMPLATE = next((p for p in [HERE / "report_template.docx",
                             HERE.parent / "reference_report" / "ESG 永續承諾驗證競賽_報告空白範本.docx"]
                 if p.exists()), None)
assert TEMPLATE, "official blank template not found (report/report_template.docx)"
OUT = HERE / "TEAM_10505_ESG永續承諾驗證競賽.docx"

GITHUB = "https://github.com/yuchengf/veripromiseesg-2026"

# ───────────────────────── content (Traditional Chinese) ─────────────────────────
SECTIONS = {
 "壹、": [
  "本隊開發與訓練環境如下。作業系統為 Ubuntu Linux（核心 6.x），程式語言為 Python 3.13。"
  "硬體採用單張 NVIDIA GeForce RTX 5090（32GB VRAM）進行所有模型訓練與推論。",
  "主要使用之套件包括：PyTorch（CUDA，建構與訓練 Transformer 模型）、HuggingFace Transformers"
  "（載入預訓練語言模型與分詞器）、scikit-learn（macro-F1 評估、交叉驗證切分、bootstrap 顯著性檢定）、"
  "NumPy 與 pandas（資料處理）、Matplotlib（結果視覺化）。",
  "預訓練模型（皆為開源、於本地端執行）：骨幹編碼器為 hfl/chinese-roberta-wwm-ext-large；"
  "kNN 標籤分布學習之文字嵌入採用 Qwen/Qwen3-Embedding-0.6B。"
  "競賽主辦單位已於官方討論串明確說明「任何訓練階段皆不限制使用開源或閉源大型語言模型」。"
  "本隊最終提交所依賴之模型全為開源且於本地推論，推論階段不呼叫任何外部 API，可完全離線重現。",
  "額外資料集：本隊曾下載並清理 ML-Promise（SemEval-2025 Task 6）多語資料用於 T3 增補實驗，"
  "惟最終未採用（跨語遷移失敗，詳見「陸」）。除此之外僅使用主辦單位提供之競賽資料集。",
  "生成式 AI 揭露：本隊於程式撰寫、實驗設計、文獻調研與數據分析過程中，"
  "使用生成式 AI 助理 Anthropic Claude（透過 Claude Code，模型為 Claude Opus 4.x，最終版本 Opus 4.8，1M context）。"
  "具體用途：協助撰寫與重構 Python 訓練／推論／評估程式碼、設計與編排消融實驗與 bootstrap 檢定、"
  "彙整相關研究文獻、以及維護實驗紀錄。競賽策略與方向、方法論紀律之堅持（抗過擬合、共同確認規則）、"
  "所有模型架構決策、實驗結論與最終提交選擇，皆由本隊人工審查與決定；致勝之核心構想"
  "（逐任務 max_length 收割、聯合結構化解碼）源自人機協作。〔請隊伍依實際情況校準本段用字與貢獻比例。〕",
 ],
 "貳、": [
  "本競賽需對企業 ESG 報告段落同時預測四項相依子任務：T1 承諾辨識 promise_status（Yes/No）、"
  "T2 證據佐證 evidence_status（Yes/No/N/A）、T3 證據品質 evidence_quality（Clear/Not Clear/Misleading/N/A）、"
  "T4 驗證時程 verification_timeline（already/within_2_years/between_2_and_5_years/more_than_5_years/N/A）。"
  "評分為加權 macro-F1：0.20×T1 + 0.30×T2 + 0.35×T3 + 0.15×T4，且僅計入解答中實際出現之類別。"
  "四任務間具邏輯硬性相依（T1=No → 下游皆 N/A；T2=No → T3=N/A），此結構為本題核心特性。",
  "整體推論流程為：段落文字 →（RoBERTa-large 編碼器，取 CLS 表徵）→ CascadeHeadV2 級聯多任務頭 →"
  " 3 seeds × 5 folds 機率平均集成 → T3 強化（NC-bias → kNN-LDL 融合 → clarity 頭覆蓋）→"
  " 聯合結構化解碼（joint decode）與 N/A 級聯規則 → 產生提交。",
  "（一）骨幹與級聯頭：採用中文 RoBERTa-large（hfl/chinese-roberta-wwm-ext-large）。"
  "CascadeHeadV2 將上游任務之 softmax 機率串接進下游頭之輸入——T1 機率送入 T2、T1 與 T2 機率送入 T3、"
  "T1 機率送入 T4——於架構層面直接編碼任務間之依賴關係。",
  "（二）逐任務損失（per-task loss）：T1 與 T2 採加權交叉熵；T3 採 Distribution-Balanced（NTR）損失並對"
  " Not Clear 類別加權（t3_nc_weight=3.0），對抗 Clear ≫ Not Clear 之嚴重不平衡；T4 採 ordinal 損失（時程具自然順序）。"
  "另採路徑損失（path loss）：T1=No 之列不計算下游損失，使下游頭僅在有意義之列學習。",
  "（三）T3 強化模組（本題最弱、權重最高之任務）：NC-bias 於推論時將 Not Clear 機率乘 exp(0.1) 再正規化；"
  "kNN-LDL 以 Qwen3-Embedding-0.6B 取 k=5 最近鄰之 T3 標籤分布，與模型機率以 0.6×model + 0.4×kNN 融合；"
  "另設一獨立之 balanced-softmax clarity 頭，於高信心（≥0.7）且非 Misleading 時硬覆蓋模型之 T3 預測。",
  "（四）規模與序列長度：骨幹為 RoBERTa-large（約 3.3 億參數），最大序列長度依任務分別設為 384 或 512"
  "（詳見「肆」與「參」）。",
  "（五）模型超參數與輸出維度：編碼器取 [CLS] 表徵（隱藏維度 1024），經 dropout（0.1）後分別送入四個任務頭，"
  "輸出維度依序為 T1=2、T2=3、T3=4、T4=5；級聯串接時採用上游之 softmax 機率分布（而非 argmax 硬標籤），"
  "以保留上游不確定性供下游頭利用，使整體可視為一個可端到端訓練之結構化多任務分類器。",
 ],
 "參、": [
  "（一）逐子任務 max_length「收割」（per-task harvest）：RoBERTa 上限為 512 token，"
  "而資料 token 長度 p95 僅 381（僅 4.8% 超過 384）。本隊訓練 @384 與 @512 兩套模型，"
  "並利用「AIdea 每次提交皆回報四個子任務分數」之特性，以單次全 @512 提交一次讀出四任務於 @512 之分數，"
  "再與已知之 @384 對照，逐任務挑選較佳來源。重大發現為：本地驗證集與真實測試集出現反轉——"
  "valid 上看似 T3@512 較佳、T4@512 較差，但 AIdea 上完全相反（T4@512 大幅勝出、T3 應留 @384）。"
  "最終來源為 T1/T3@384、T2/T4@512。此發現體現「本地 CV 與測試分布不一致時，應以同分布訊號為準」。",
  "（二）聯合結構化解碼（joint / structured decoding）—— 核心貢獻：原本解碼為貪婪、單向"
  "（先 argmax 上游 gate，T1=No 即強制下游為 N/A），僅用到依賴之前向。本隊改為對每列列舉三個合法聯合分支"
  "（No / Yes-No / Yes-Yes），選擇使四任務聯合對數機率最大者，並使用模型自身對下游 N/A 之機率，"
  "使下游之高信心可反向修正上游 gate。其性質為推論期、不需重訓、且具結構性（用真實依賴與模型自身機率，"
  "而非調參捷徑）。效果為「同一批機率、僅改解碼」即在 AIdea 上四個子任務同時提升，總分 0.6232→0.6249，"
  "並於 Private 重排中持續轉移。",
  "（三）kNN 標籤分布學習與獨立 clarity 頭強化 T3：以外部嵌入模型之鄰居標籤分布與獨立分類頭，"
  "對權重最高、最難之 T3 任務提供與骨幹去相關之第二意見。",
  "（四）抗過擬合之方法論紀律：本隊將 CV、Public LB、Private 視為三個有偏之噪音估計，"
  "僅採用能被兩個獨立訊號共同確認（convergent evidence）、且具結構性之改動；對只在單一榜單變好之改動一律視為過擬合嫌疑而拒絕。"
  "此紀律是 Private 名次上升之主因（詳見「陸」）。",
 ],
 "肆、": [
  "競賽資料集 VeriPromiseESG4K 取自台灣 50 指數成分股之 ESG 報告書，每列為一段文字並附四任務標註。"
  "本隊使用兩套資料設定：final_data（1601 筆，含本地 valid 399 筆，用於選型與消融）與"
  " retrain_data（2000 筆，用於產生最終提交之模型）。",
  "（一）N/A 級聯結構之驗證：本隊以訓練資料實際統計驗證級聯規則完全成立——T1=No 時 T2/T3/T4 必為 N/A（100%）；"
  "T1=Yes 且 T2=No 時 T3 必為 N/A（100%）但 T4 仍為時程類別（0% 為 N/A）；T1=Yes 且 T2=Yes 時 T3/T4 皆為實類別。"
  "推論時據此套用 N/A 級聯規則，並確認最終提交無級聯違規列。",
  "（二）序列長度與截斷分析：token 長度分布為 p50=174、p90=319、p95=381、p99=550。"
  "據此將 max_length 設為 384（覆蓋 95% 之資料）為主，並另訓練 @512 版本供逐任務收割比較（見「參」）。",
  "（三）稀有類別處理：T3 之 Not Clear、Misleading 與 T4 之 within_2_years 皆為稀有類別。"
  "訓練時以 augment_rare 對稀有類別過採樣，並於 T3 損失加權 Not Clear；推論時以 NC-bias 與 kNN-LDL 進一步調整。"
  "其中 Misleading 經驗證為無法由文字內容判定之主觀標註（見「陸」），故本隊不對其強行預測，以免注入假陽性。",
  "（四）文字前處理：本隊以保留語意完整性為原則，僅去除段落前後空白與多餘之連續空白字元，"
  "不額外進行中文斷詞（交由 RoBERTa 之子詞分詞器處理），亦不移除標點與英數、不做大小寫正規化，"
  "以避免破壞 ESG 文句中關鍵之數字、年限、比率與專有名詞；分詞後超過設定長度之序列以尾端截斷處理。",
 ],
 "伍、": [
  "（一）訓練配方：主力配方（代號 FC1R）於 CascadeHeadV2 上加入 R-Drop（α=0.5）與 SWA（自第 7 個 epoch 起），"
  "並開啟 per-task loss、augment_rare、deep_cascade 與 t3_nc_weight=3.0。"
  "本隊亦訓練 base、SharpReCL、MR2 等其他配方供集成消融，惟最終以 FC1R 轉移最佳而採用。"
  "各配方共用之基礎超參數為：AdamW 最佳化器、學習率 2e-5 並搭配線性 warmup 學習率排程、"
  "每折最多訓練 10 個 epoch、採早停（patience=3）、並以驗證集之加權 macro-F1 選取最佳檢查點儲存；"
  "批次大小於 R-Drop 配方設為 8、其餘配方為 16，記憶體不足時自動退回較小批次重試。",
  "（二）最佳化與集成：採 AdamW 最佳化器；以 5-fold 交叉驗證 × 3 個隨機種子（42/123/456）做機率平均集成。"
  "種子曲線（1→2→3 seed：0.658→0.673→0.679）顯示報酬遞減，3 seed 為最佳折衷；"
  "本隊驗證 5 seed 相對 3 seed 之增益在噪音範圍內，故不採用。",
  "（三）兩階段資料流程：先於 final_data（1601）訓練、以本地 valid 399 選型與做消融；"
  "確立關鍵設計（per-task @384/@512、kNN α=0.4、NC-bias 0.1、clarity 門檻 0.7）後，"
  "再於 retrain_data（2000）自頭重訓 RT 模型以最大化可用資訊，並據此生成最終提交。",
  "（四）提交紀律：競賽每日提交配額有限（組內共用）。本隊維護提交紀錄，對每次提交記錄"
  "「預期分數、欲驗證之假設、事後結論」，並盡量以單一變因為原則，使 LB 回饋能明確歸因；"
  "對昂貴之多種子重訓，先以最便宜之隔離測試（1-seed 對 1-seed）確認方法有效後才放大，避免浪費配額與運算。",
  "（五）最終提交生成：以 retrain 模型產生各任務機率（gen_rt_more.py），融合 kNN-LDL 與 clarity，"
  "最後以聯合結構化解碼（gen_joint_aidea.py，wgate=2.0）產生 aidea_joint_w2.0.csv，"
  "並逐項檢查列數（2000）、標籤字串、N/A 級聯一致性後方才提交。",
 ],
 "陸、": [
  "本隊 Public LB 成績逐步演進：0.6188（clarity-on-3way）→ 0.6221（pure @512）→ 0.6232"
  "（逐任務 @384/@512 混合）→ 0.6249（聯合結構化解碼），每一步皆為結構性且經 AIdea 確認之增益。"
  "最終 Public 為 0.6249259（第 15 名）。",
  "最關鍵之驗證在 Private 重排：本隊提交於 Private LB 取得 0.6432854（第 12 名），"
  "相對 Public 分數提升約 +0.018、名次上升 3 名；而許多在截止前大量刷 Public 的隊伍則退步。"
  "此結果直接印證本隊「只押結構性、共同確認之改動，不追逐單一榜單」之方法論——"
  "結構性、未過度擬合之提交，正是在 Private 重排中得以保持並上升的關鍵。",
  "最終提交之穩健性驗證：對解碼權重 wgate∈[1.0,2.5] 擾動時，提交僅改變 0.1–0.2% 之儲存格（非刀鋒解）；"
  "最終檔有 98.2% 與已確認之基底相同（為「已確認基底＋微小結構性修正」，而非脆弱新建構）；"
  "格式與 N/A 級聯經驗證無誤；且為結構性提交而非截止前刷榜，故能穩健面對分布變動。",
  "天花板診斷——T3 為 label-limited：本隊從六個獨立角度（14B 大型語言模型判別、ESG 領域特徵 re-ranker、"
  "跨 backbone 集成、kNN、事後 logit 調整、訓練期條件鏈式 aux）以及 SemEval-2025 Task 6 與 LeWiDi 之文獻，"
  "確認 T3 之 Clear↔Not Clear↔Misleading 邊界無法由文字內容判定（Misleading 屬主觀標註，訊號不在輸入中）；"
  "SemEval 同題之頂尖系統亦僅約 0.52–0.53。因此本隊停止於 T3 投入，將資源集中於可轉移之結構性改動。",
  "失敗案例與否決之方法（誠實的負面結果）：company/ticker prior 融合（valid 顯著正但 AIdea 過擬合退步）、"
  "12-way 跨配方集成（valid 0.681 最高卻 AIdea 0.614 最差，為 valid→test 反轉鐵證）、"
  "跨 backbone 集成（XLM-R/mmBERT/BGE-M3/DeBERTa 皆弱於 RoBERTa）、Misleading 與 gate 之 LLM 判別、"
  "領域特徵 re-ranker、SupCon 對比微調、ordinal/OLL、條件鏈式 aux、post-hoc logit 調整與 per-class threshold、"
  "soft-macro-F1、半監督/pseudo-label——皆因過擬合或碰到 T3 之 label-limit 而否決。"
  "統一結論：上述多為「重塑輸出/損失」之方法，當瓶頸是輸入訊號/標籤本身之限制時，無法補回不存在之訊號。",
  "研究限制與未來方向：（一）受限於時程與配額，部分方向（如以外部資料將 Misleading 連結至已知漂綠案例）"
  "未能完整驗證；（二）若一開始即釐清本地 CV 與測試分布不一致，可更早採用同分布訊號為準；"
  "（三）真正能突破 T3 天花板者，須為新的輸入訊號（外部資料或更強表徵），而非輸出端方法。"
  "整體而言，本隊在嚴格之抗過擬合紀律下，由 0.6188 逐步提升至最終 Private 0.6432854（第 12 名），"
  "驗證了任務拆分、依賴建模（聯合解碼）、與紀律性決策之有效性。",
 ],
 "捌、": [
  "外部資源：開源工具包含 PyTorch、HuggingFace Transformers、scikit-learn、NumPy、pandas、Matplotlib（皆為公開套件）；"
  "開源預訓練模型 hfl/chinese-roberta-wwm-ext-large 與 Qwen/Qwen3-Embedding-0.6B（皆於本地執行）；"
  "並使用生成式 AI 助理 Anthropic Claude（Claude Code，Opus 4.x／4.8）協助程式撰寫與分析（已於「壹、環境」揭露）。"
  "資料方面除主辦單位競賽資料集外，曾於實驗中使用 ML-Promise（SemEval-2025 Task 6）資料但最終未採用。",
 ],
}

CODE_LINK = f"GitHub 連結：{GITHUB}（公開倉庫，可直接開啟瀏覽與 clone 執行）"
CODE_CONTENTS = (
  "程式碼倉庫包含：模型／級聯頭／損失／訓練／k-fold／推論之主程式 esg_main.py；"
  "以 retrain 模型產生提交與 kNN-LDL 融合之 gen_rt_more.py；聯合結構化解碼之 gen_joint_aidea.py；"
  "獨立 clarity 頭之 clarity_head.py；本地 valid 評估與機率快取之 eval_single_run.py；"
  "並附 README.md 詳盡說明配置環境安裝、資料路徑、重要模塊輸入／輸出、訓練指令與重現步驟"
  "（最終提交可由 python gen_joint_aidea.py 2.0 重現 aidea_joint_w2.0.csv），以利第三方用戶除錯、重新訓練與重現結果。")

REFS = [
 "Cui, Y., Che, W., Liu, T., Qin, B., & Yang, Z. (2021). Pre-training with whole word masking for Chinese BERT. IEEE/ACM Transactions on Audio, Speech, and Language Processing, 29, 3504–3514.",
 "Liu, Y., Ott, M., Goyal, N., Du, J., Joshi, M., Chen, D., Levy, O., Lewis, M., Zettlemoyer, L., & Stoyanov, V. (2019). RoBERTa: A robustly optimized BERT pretraining approach. arXiv:1907.11692.",
 "Vaswani, A., Shazeer, N., Parmar, N., Uszkoreit, J., Jones, L., Gomez, A. N., Kaiser, Ł., & Polosukhin, I. (2017). Attention is all you need. In Advances in Neural Information Processing Systems (NeurIPS).",
 "Lin, T.-Y., Goyal, P., Girshick, R., He, K., & Dollár, P. (2017). Focal loss for dense object detection. In Proceedings of the IEEE International Conference on Computer Vision (ICCV).",
 "Wu, T., Huang, Q., Liu, Z., Wang, Y., & Lin, D. (2020). Distribution-balanced loss for multi-label classification in long-tailed datasets. In Proceedings of the European Conference on Computer Vision (ECCV).",
 "Wu, L., Li, J., Wang, Y., Meng, Q., Qin, T., Chen, W., Zhang, M., Liu, T.-Y. (2021). R-Drop: Regularized dropout for neural networks. In Advances in Neural Information Processing Systems (NeurIPS).",
 "Menon, A. K., Jayasumana, S., Rawat, A. S., Jain, H., Veit, A., & Kumar, S. (2021). Long-tail learning via logit adjustment. In International Conference on Learning Representations (ICLR).",
 "Shi, X., Cao, W., & Raschka, S. (2023). Deep neural networks for rank-consistent ordinal regression based on conditional probabilities (CORN). Pattern Analysis and Applications, 26, 941–955.",
 "Bénédict, G., Koops, V., Odijk, D., & de Rijke, M. (2022). sigmoidF1: A smooth F1 score surrogate loss for multilabel classification. Transactions on Machine Learning Research.",
 "Ren, J., Yu, C., Sheng, S., Ma, X., Zhao, H., Yi, S., & Li, H. (2020). Balanced meta-softmax for long-tailed visual recognition. In Advances in Neural Information Processing Systems (NeurIPS).",
 "SemEval-2025 Task 6 organizers. (2025). SemEval-2025 Task 6: Multinational, multilingual, multi-industry promise verification (PromiseEval). In Proceedings of the 19th International Workshop on Semantic Evaluation (SemEval-2025).",
 "Leonardelli, E., et al. (2025). LeWiDi-2025: The third edition of the Learning With Disagreements shared task. arXiv:2510.08460.",
]

# ───────────────────────── helpers ─────────────────────────
def set_run_font(r, size=12, bold=False):
    r.font.name = EN_FONT; r.font.size = Pt(size); r.font.bold = bold
    rpr = r._element.get_or_add_rPr()
    rf = rpr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts'); rpr.append(rf)
    rf.set(qn('w:ascii'), EN_FONT); rf.set(qn('w:hAnsi'), EN_FONT); rf.set(qn('w:eastAsia'), CN_FONT)


def style_para(p):
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.15
    if pf.space_after is None:
        pf.space_after = Pt(6)


def write_para(p, text, size=12, bold=False, hang=False):
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    style_para(p)
    if hang:
        p.paragraph_format.left_indent = Cm(0.74)
        p.paragraph_format.first_line_indent = Cm(-0.74)
    set_run_font(p.add_run(text), size, bold)


def new_para_after(anchor):
    el = OxmlElement('w:p')
    anchor._p.addnext(el)
    return Paragraph(el, anchor._parent)


def fill_block(placeholder, paragraphs):
    """Replace a single [內文] placeholder with N body paragraphs; return last paragraph."""
    write_para(placeholder, paragraphs[0])
    cur = placeholder
    for txt in paragraphs[1:]:
        cur = new_para_after(cur); write_para(cur, txt)
    return cur


def set_cell(cell, text):
    p = cell.paragraphs[0]
    write_para(p, text)


# ───────────────────────── build ─────────────────────────
doc = Document(str(TEMPLATE))
HEAD_KEYS = list(SECTIONS.keys()) + ["柒、"]

current = None
for p in list(doc.paragraphs):
    t = p.text.strip()
    # drop the「（空白範本）」subtitle
    if t == "（空白範本）":
        p._element.getparent().remove(p._element); continue
    # team-info fields
    if t.startswith("隊伍："):
        write_para(p, "隊伍：TEAM_10505（學生組）"); continue
    if t.startswith("Private leaderboard"):
        write_para(p, "Private leaderboard：0.6432854 / Rank 12"); continue
    # track current section heading
    for k in HEAD_KEYS:
        if t.startswith(k):
            current = k; break
    # section bodies
    if t == "[內文]" and current in SECTIONS:
        last = fill_block(p, SECTIONS[current])
        if current == "捌、":   # append references after the external-resources paragraph
            last = new_para_after(last); write_para(last, "參考文獻（APA 格式）：")
            for ref in REFS:
                last = new_para_after(last); write_para(last, ref, hang=True)
        current = None
    elif t.startswith("GitHub 連結") and current == "柒、":
        last = fill_block(p, [CODE_LINK, CODE_CONTENTS]); current = None

# author-contact table: fill team name / private score / rank (member rows left blank)
tbl = doc.tables[0]
VALS = {"隊伍名稱": "TEAM_10505",
        "Private Leaderboard 成績": "0.6432854",
        "Private Leaderboard 名次": "12"}
cells = tbl.rows[0].cells
for i, c in enumerate(cells):
    if c.text.strip() in VALS and i + 1 < len(cells):
        set_cell(cells[i + 1], VALS[c.text.strip()])

# enforce 1.15 line spacing on every paragraph (incl. headings + table cells)
for p in doc.paragraphs:
    style_para(p)
for tb in doc.tables:
    for row in tb.rows:
        for c in row.cells:
            for p in c.paragraphs:
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                p.paragraph_format.line_spacing = 1.15

doc.save(str(OUT))
print("Saved:", OUT.name)
